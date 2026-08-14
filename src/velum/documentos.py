"""Lectura y escritura de documentos.

Nada se reconstruye. El DOCX se abre, se sustituye dentro de sus propios
elementos y se guarda: estilos, tablas, numeración, cabeceras, pies y notas
salen como entraron.
"""

from __future__ import annotations

import hashlib
import shutil
from dataclasses import dataclass, field
from pathlib import Path

EXTENSIONES_TEXTO = {".txt", ".md", ".markdown", ".rst"}
EXTENSIONES_DOCX = {".docx"}
EXTENSIONES_SOPORTADAS = EXTENSIONES_TEXTO | EXTENSIONES_DOCX

# Formatos que se detectan para poder avisar, pero que esta versión no reescribe.
EXTENSIONES_NO_SOPORTADAS = {".pdf", ".doc", ".rtf", ".odt", ".pages"}


class FormatoNoSoportado(RuntimeError):
    pass


@dataclass
class Documento:
    """Un documento abierto como lista de bloques de texto sustituibles."""

    ruta: Path
    formato: str
    bloques: list[str]
    partes: list[str] = field(default_factory=list)   # de dónde viene cada bloque
    _docx: object | None = None
    _referencias: list[object] = field(default_factory=list)

    @property
    def texto_plano(self) -> str:
        return "\n".join(self.bloques)


def hash_fichero(ruta: Path) -> str:
    resumen = hashlib.sha256()
    with open(ruta, "rb") as fichero:
        for trozo in iter(lambda: fichero.read(65536), b""):
            resumen.update(trozo)
    return resumen.hexdigest()


# ---------------------------------------------------------------------------
# Texto plano
# ---------------------------------------------------------------------------

def _abrir_texto(ruta: Path) -> Documento:
    contenido = ruta.read_text(encoding="utf-8", errors="replace")
    lineas = contenido.split("\n")
    return Documento(ruta=ruta, formato="texto", bloques=lineas,
                     partes=["cuerpo"] * len(lineas))


def _guardar_texto(documento: Documento, destino: Path) -> None:
    destino.write_text("\n".join(documento.bloques), encoding="utf-8")


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

def _parrafos_docx(documento) -> list[tuple[object, str]]:
    """Todos los párrafos del documento, con la parte de la que proceden."""
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    recogidos: list[tuple[object, str]] = []

    def recorrer_tabla(tabla: Table, parte: str) -> None:
        for fila in tabla.rows:
            for celda in fila.cells:
                for parrafo in celda.paragraphs:
                    recogidos.append((parrafo, parte))
                for anidada in celda.tables:
                    recorrer_tabla(anidada, parte)

    for parrafo in documento.paragraphs:
        recogidos.append((parrafo, "cuerpo"))
    for tabla in documento.tables:
        recorrer_tabla(tabla, "tabla")

    for seccion in documento.sections:
        for contenedor, nombre in (
            (seccion.header, "cabecera"),
            (seccion.footer, "pie"),
            (seccion.first_page_header, "cabecera"),
            (seccion.first_page_footer, "pie"),
            (seccion.even_page_header, "cabecera"),
            (seccion.even_page_footer, "pie"),
        ):
            if contenedor is None:
                continue
            for parrafo in contenedor.paragraphs:
                recogidos.append((parrafo, nombre))
            for tabla in contenedor.tables:
                recorrer_tabla(tabla, nombre)

    # Deduplica párrafos compartidos entre secciones enlazadas.
    vistos: set[int] = set()
    unicos: list[tuple[object, str]] = []
    for parrafo, parte in recogidos:
        identidad = id(parrafo._p)
        if identidad in vistos:
            continue
        vistos.add(identidad)
        unicos.append((parrafo, parte))
    return unicos


def _nodos_notas(documento) -> list[tuple[object, str]]:
    """Nodos <w:t> de notas al pie y notas finales.

    Limitación honesta: aquí se sustituye nodo a nodo, de modo que un dato
    partido entre dos nodos de una nota al pie puede escapar. El control de
    calidad final lo detecta y se avisa en el acta.
    """
    from docx.oxml.ns import qn

    nodos: list[tuple[object, str]] = []
    paquete = documento.part.package
    for parte in paquete.parts:
        nombre = str(parte.partname)
        if not (nombre.endswith("footnotes.xml") or nombre.endswith("endnotes.xml")):
            continue
        elemento = getattr(parte, "element", None)
        if elemento is None:
            continue
        etiqueta = "nota" if "footnotes" in nombre else "nota final"
        for nodo in elemento.iter(qn("w:t")):
            nodos.append((nodo, etiqueta))
    return nodos


def _abrir_docx(ruta: Path) -> Documento:
    try:
        import docx  # type: ignore
    except ImportError as error:  # pragma: no cover
        raise FormatoNoSoportado(
            "Para procesar .docx hace falta python-docx: pip install python-docx"
        ) from error

    documento_docx = docx.Document(str(ruta))
    referencias: list[object] = []
    bloques: list[str] = []
    partes: list[str] = []

    for parrafo, parte in _parrafos_docx(documento_docx):
        referencias.append(("parrafo", parrafo))
        bloques.append(parrafo.text)
        partes.append(parte)

    for nodo, parte in _nodos_notas(documento_docx):
        referencias.append(("nodo", nodo))
        bloques.append(nodo.text or "")
        partes.append(parte)

    return Documento(
        ruta=ruta, formato="docx", bloques=bloques, partes=partes,
        _docx=documento_docx, _referencias=referencias,
    )


def _sustituir_en_parrafo(parrafo, sustituciones: list[tuple[int, int, str]]) -> None:
    """Aplica sustituciones respetando el formato de cada run."""
    runs = parrafo.runs
    if not runs or not sustituciones:
        return

    def desplazamientos() -> list[tuple[int, int]]:
        posicion = 0
        tramos = []
        for run in runs:
            longitud = len(run.text or "")
            tramos.append((posicion, posicion + longitud))
            posicion += longitud
        return tramos

    tramos = desplazamientos()

    for inicio, fin, marcador in sorted(sustituciones, reverse=True):
        afectados = [i for i, (a, b) in enumerate(tramos) if a < fin and b > inicio]
        if not afectados:
            continue

        primero, ultimo = afectados[0], afectados[-1]
        inicio_primero = tramos[primero][0]
        inicio_ultimo = tramos[ultimo][0]

        prefijo = (runs[primero].text or "")[: inicio - inicio_primero]
        sufijo = (runs[ultimo].text or "")[fin - inicio_ultimo :]

        if primero == ultimo:
            runs[primero].text = prefijo + marcador + sufijo
        else:
            runs[primero].text = prefijo + marcador
            for indice in afectados[1:-1]:
                runs[indice].text = ""
            runs[ultimo].text = sufijo

        tramos = desplazamientos()


def _guardar_docx(documento: Documento, destino: Path,
                  sustituciones: dict[int, list[tuple[int, int, str]]]) -> None:
    for indice, (clase, referencia) in enumerate(documento._referencias):
        cambios = sustituciones.get(indice)
        if not cambios:
            continue
        if clase == "parrafo":
            _sustituir_en_parrafo(referencia, cambios)
        else:  # nodo <w:t> de una nota
            from .motor import Anonimizador

            referencia.text = Anonimizador.aplicar_a_texto(referencia.text or "", cambios)

    documento._docx.save(str(destino))  # type: ignore[union-attr]


# ---------------------------------------------------------------------------
# API pública
# ---------------------------------------------------------------------------

def abrir(ruta: str | Path) -> Documento:
    camino = Path(ruta).expanduser()
    if not camino.exists():
        raise FileNotFoundError(f"No existe el fichero: {camino}")

    extension = camino.suffix.lower()
    if extension in EXTENSIONES_DOCX:
        return _abrir_docx(camino)
    if extension in EXTENSIONES_TEXTO:
        return _abrir_texto(camino)
    if extension in EXTENSIONES_NO_SOPORTADAS:
        raise FormatoNoSoportado(
            f"VELUM aún no reescribe ficheros {extension}. "
            "Formatos admitidos en esta versión: .docx, .txt, .md. "
            "Convierta el documento a .docx y vuelva a intentarlo."
        )
    raise FormatoNoSoportado(f"Extensión no reconocida: {extension}")


def guardar(
    documento: Documento,
    destino: str | Path,
    sustituciones: dict[int, list[tuple[int, int, str]]] | None = None,
) -> Path:
    camino = Path(destino).expanduser()
    camino.parent.mkdir(parents=True, exist_ok=True)

    if documento.formato == "docx":
        _guardar_docx(documento, camino, sustituciones or {})
    else:
        _guardar_texto(documento, camino)
    return camino


def ruta_de_salida(origen: Path, carpeta_salida: Path | None, sufijo: str = "_anonimizado") -> Path:
    destino_dir = carpeta_salida or origen.parent
    return Path(destino_dir) / f"{origen.stem}{sufijo}{origen.suffix}"


def listar_carpeta(carpeta: str | Path, recursivo: bool = True) -> tuple[list[Path], list[Path]]:
    """Devuelve (compatibles, no compatibles) sin abrir ningún fichero."""
    raiz = Path(carpeta).expanduser()
    if not raiz.is_dir():
        raise NotADirectoryError(f"No es una carpeta: {raiz}")

    patron = "**/*" if recursivo else "*"
    compatibles: list[Path] = []
    incompatibles: list[Path] = []

    for camino in sorted(raiz.glob(patron)):
        if not camino.is_file() or camino.name.startswith("~$"):
            continue
        if "_anonimizado" in camino.stem:
            continue
        extension = camino.suffix.lower()
        if extension in EXTENSIONES_SOPORTADAS:
            compatibles.append(camino)
        elif extension in EXTENSIONES_NO_SOPORTADAS:
            incompatibles.append(camino)
    return compatibles, incompatibles


def copia_de_seguridad(origen: Path) -> Path:
    destino = origen.with_suffix(origen.suffix + ".bak")
    if not destino.exists():
        shutil.copy2(origen, destino)
    return destino
